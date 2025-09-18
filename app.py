import os
from flask import Flask, request, render_template_string, send_file, redirect, url_for, flash, jsonify
from werkzeug.utils import secure_filename
from transformers import pipeline, MarianMTModel, MarianTokenizer
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime
from groq import Groq
from docx import Document
from PyPDF2 import PdfReader

from dotenv import load_dotenv
import threading
import contextlib
import torch
from docx.shared import RGBColor

# Load .env
load_dotenv()

# Flask app
app = Flask(__name__)
app.secret_key = "supersecretkey"

# Folders
UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "outputs"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# Allowed extensions
ALLOWED_EXTENSIONS = {"txt", "docx", "pdf"}

# Groq client
# Groq client - Add your API key here
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
if not GROQ_API_KEY:
    raise ValueError("GROQ_API_KEY environment variable is required")
groq_client = Groq(api_key=GROQ_API_KEY)

# Thread-safe model cache
model_cache = {}
model_lock = threading.Lock()

# HTML Template
HTML_TEMPLATE = """ 
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Advanced AI Translator</title>
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body { font-family: 'Segoe UI', sans-serif; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: #333; min-height: 100vh; }
        .container { max-width: 1000px; margin: 0 auto; padding: 20px; }
        .header { text-align: center; color: white; margin-bottom: 30px; }
        .header h1 { font-size: 2.5em; text-shadow: 2px 2px 4px rgba(0,0,0,0.3); margin-bottom: 10px; }
        .card { background: white; border-radius: 15px; padding: 30px; margin-bottom: 20px; box-shadow: 0 10px 30px rgba(0,0,0,0.2); }
        textarea { width: 100%; padding: 15px; border: 2px solid #e9ecef; border-radius: 8px; font-size: 16px; resize: vertical; min-height: 120px; font-family: inherit; }
        textarea:focus { outline: none; border-color: #667eea; box-shadow: 0 0 0 3px rgba(102,126,234,0.1); }
        .flash { padding: 15px 20px; margin: 15px 0; border-radius: 8px; font-weight: 500; }
        .error { background: #f8d7da; border: 1px solid #f5c6cb; color: #721c24; }
        .success { background: #d4edda; border: 1px solid #c3e6cb; color: #155724; }
        #translatedOutput { white-space: pre-wrap; border: 2px solid #e9ecef; padding: 20px; min-height: 120px; border-radius: 8px; background: #f8f9fa; font-family: inherit; line-height: 1.5; }
        .btn { background: #667eea; color: white; padding: 12px 24px; border: none; border-radius: 8px; cursor: pointer; font-size: 16px; font-weight: 600; transition: all 0.3s ease; }
        .btn:hover { background: #5a67d8; transform: translateY(-2px); }
        .file-upload { border: 3px dashed #dee2e6; border-radius: 10px; padding: 20px; text-align: center; background: #f8f9fa; margin: 15px 0; transition: all 0.3s ease; }
        .file-upload:hover { border-color: #667eea; background: #e3f2fd; }
        .language-info { background: #e3f2fd; border: 1px solid #90caf9; border-radius: 8px; padding: 15px; margin: 15px 0; color: #0d47a1; font-weight: 500; }
        .processing-indicator { display: none; text-align: center; color: #667eea; font-style: italic; margin: 10px 0; }
        .tabs { display: flex; margin-bottom: 20px; border-radius: 10px; overflow: hidden; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }
        .tab { flex: 1; padding: 15px; background: #f8f9fa; border: none; cursor: pointer; font-size: 16px; transition: all 0.3s ease; }
        .tab.active { background: #667eea; color: white; }
        .tab:hover:not(.active) { background: #e9ecef; }
        .tab-content { display: none; }
        .tab-content.active { display: block; }
        @media (max-width: 768px) { .container { padding: 10px; } .header h1 { font-size: 2em; } .tabs { flex-direction: column; } }
    </style>
</head>
<body>
<div class="container">
    <div class="header">
        <h1>🤖 Advanced AI Translator</h1>
        <p>Powered by Groq AI Detection & Helsinki-NLP Models</p>
    </div>

    {% with messages = get_flashed_messages(with_categories=true) %}
      {% if messages %}
        {% for category, message in messages %}
          <div class="flash {{ category }}">{{ message }}</div>
        {% endfor %}
      {% endif %}
    {% endwith %}

    <div class="card">
        <div class="tabs">
            <button class="tab active" onclick="switchTab(event, 'live-tab')">💬 Live Translation</button>
            <button class="tab" onclick="switchTab(event, 'file-tab')">📄 File Translation</button>
        </div>

        <div id="live-tab" class="tab-content active">
            <h3>🔄 Real-time Translation</h3>
            <div class="language-info">
                <strong>🎯 Target:</strong> English (Auto-detected source language)
            </div>
            <textarea id="inputText" placeholder="Type or paste your text here in any language...">{{ text or "" }}</textarea>
            <div class="processing-indicator" id="processing">🔄 Translating...</div>
            <h4>✅ Translation Result</h4>
            <div id="translatedOutput"></div>
        </div>

        <div id="file-tab" class="tab-content">
            <form method="POST" enctype="multipart/form-data">
                <h3>📁 Document Translation</h3>
                <div class="language-info">
                    <strong>🎯 Target:</strong> English | <strong>🔍 Source:</strong> Auto-detected
                </div>
                <div class="file-upload">
                    <input type="file" name="file" accept=".txt,.docx,.pdf" id="fileInput">
                    <p>📁 Choose a file to translate (TXT, DOCX, PDF)</p>
                    <small>Maximum file size: 16MB</small>
                </div>
                <input type="hidden" name="input_type" value="file">
                <button type="submit" class="btn">🚀 Translate Document</button>
            </form>
        </div>
    </div>

    {% if translated_text %}
    <div class="card">
        <h3>📄 Translated Document Preview</h3>
        {% if detected_lang %}
        <div class="language-info">
            <strong>🔍 Detected Language:</strong> {{ detected_lang.upper() }} → <strong>🎯 English</strong>
        </div>
        {% endif %}
        <textarea readonly style="min-height: 200px;">{{ translated_text }}</textarea>
        {% if download_url %}
        <div style="margin-top: 15px;">
            <a href="{{ download_url }}" class="btn">💾 Download Full Document</a>
        </div>
        {% endif %}
    </div>
    {% endif %}

    {% if processing_info %}
    <div class="language-info">{{ processing_info }}</div>
    {% endif %}
</div>

<script>
const inputBox = document.getElementById("inputText");
const outputBox = document.getElementById("translatedOutput");
const processingIndicator = document.getElementById("processing");
let timeout = null;
let lastTranslation = "";

function switchTab(evt, tabName) {
    var i, tabcontent, tabs;
    tabcontent = document.getElementsByClassName("tab-content");
    for (i = 0; i < tabcontent.length; i++) tabcontent[i].classList.remove("active");
    tabs = document.getElementsByClassName("tab");
    for (i = 0; i < tabs.length; i++) tabs[i].classList.remove("active");
    document.getElementById(tabName).classList.add("active");
    evt.currentTarget.classList.add("active");
}

inputBox.addEventListener("input", function () {
    clearTimeout(timeout);
    const text = inputBox.value;
    if (!text.trim()) {
        outputBox.innerText = " ";
        lastTranslation = "";
        processingIndicator.style.display = "none";
        return;
    }
    processingIndicator.style.display = "block";
    timeout = setTimeout(() => {
        fetch("/ajax_translate", {
            method: "POST",
            headers: { "Content-Type": "application/json" },
            body: JSON.stringify({ text: text })
        })
        .then(res => res.json())
        .then(data => {
            processingIndicator.style.display = "none";
            if (data.translation) {
                const newTranslation = data.translation;
                if (newTranslation.startsWith(lastTranslation)) {
                    outputBox.innerText += newTranslation.slice(lastTranslation.length);
                } else {
                    outputBox.innerText = `[${data.detected_lang.toUpperCase()} → EN]\n\n${newTranslation}`;
                }
                lastTranslation = newTranslation;
            } else outputBox.innerText = "[Translation unavailable]";
        })
        .catch(err => {
            processingIndicator.style.display = "none";
            outputBox.innerText = "[Translation error]";
            console.error(err);
        });
    }, 800);
});

document.getElementById('fileInput').addEventListener('change', function(e) {
    const fileName = e.target.files[0]?.name || '';
    const uploadDiv = document.querySelector('.file-upload p');
    if (fileName) uploadDiv.innerHTML = `📁 Selected: <strong>${fileName}</strong><br><small>Ready to translate!</small>`;
});
</script>
</body>
</html>

"""

def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

# ------------------------------
# Language Detection with Dialects
# ------------------------------
def detect_language_with_global_dialects(text):
    """Detect language with ISO 639-1 + dialect"""
    try:
        sample_text = text[:300].strip()
        response = groq_client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[
                {
                    "role": "system",
                    "content": """You are a world language and dialect detection expert. 
Respond ONLY with ISO 639-1 + region variant (e.g., 'en-US', 'pt-BR', 'es-MX')."""
                },
                {"role": "user", "content": f"Detect language and dialect: {sample_text}"}
            ],
            temperature=0.1,
            max_tokens=10
        )
        detected = response.choices[0].message.content.strip()
        lang_code = ''.join(c for c in detected if c.isalnum() or c == '-')
        valid_codes = {
            # English
            'en-US', 'en-GB', 'en-CA', 'en-AU', 'en-IN', 'en-NZ', 'en-ZA', 'en-IE', 'en-SG',
            # Portuguese
            'pt-BR', 'pt-PT', 'pt-AO', 'pt-MZ', 'pt-CV', 'pt-GW',
            # Spanish
            'es-ES', 'es-MX', 'es-AR', 'es-CO', 'es-CL', 'es-PE', 'es-VE', 'es-EC', 'es-UY', 'es-BO', 'es-PR', 'es-DO',
            # French
            'fr-FR', 'fr-CA', 'fr-BE', 'fr-CH', 'fr-LU', 'fr-MA', 'fr-DZ', 'fr-TN', 'fr-SN', 'fr-CI',
            # German
            'de-DE', 'de-CH', 'de-AT', 'de-LU', 'de-BE',
            # Chinese
            'zh-CN', 'zh-TW', 'zh-HK', 'zh-SG', 'zh-MO',
            # Arabic
            'ar-SA', 'ar-EG', 'ar-AE', 'ar-MA', 'ar-LB', 'ar-SY', 'ar-IQ', 'ar-JO', 'ar-KW', 'ar-OM', 'ar-QA',
            # Other widely used
            'hi-IN', 'ja-JP', 'ko-KR', 'it-IT', 'it-CH', 'nl-NL', 'nl-BE', 'nl-SR', 'ru-RU', 'pl-PL', 'cs-CZ', 
            'sk-SK', 'hr-HR', 'sl-SI', 'hu-HU', 'ro-RO', 'bg-BG', 'et-EE', 'lv-LV', 'lt-LT',
            'sv-SE', 'sv-FI', 'no-NO', 'da-DK', 'fi-FI', 'is-IS',
            'sw-KE', 'sw-TZ', 'af-ZA', 'zu-ZA', 'xh-ZA', 'mt-MT', 'ga-IE', 'cy-GB',
            'tr-TR', 'vi-VN', 'th-TH', 'id-ID', 'ms-MY', 'fil-PH', 'bn-BD', 'pa-IN', 'ur-PK'
        }
        if lang_code in valid_codes:
            return lang_code
        else:
            return "en-US"
    except Exception as e:
        print(f"[Language Detection Error]: {e}")
        return "en-US"

# ------------------------------
# Marian Model Loader & Translator
# ------------------------------
def get_marian_model(src_lang, tgt_lang="en"):
    """Get or load MarianMT model with caching"""
    model_key = f"{src_lang}-{tgt_lang}"
    with model_lock:
        if model_key in model_cache:
            return model_cache[model_key]

        # Map dialect to base language for Helsinki-NLP
        base_src = src_lang.split("-")[0].lower()
        model_name = f"Helsinki-NLP/opus-mt-{base_src}-{tgt_lang}"

        try:
            tokenizer = MarianTokenizer.from_pretrained(model_name)
            model = MarianMTModel.from_pretrained(model_name)
            model_cache[model_key] = {"tokenizer": tokenizer, "model": model}
            return model_cache[model_key]
        except Exception as e:
            # Fallback multilingual model
            print(f"[Model Load Failed]: {model_name}, {e}")
            translator = pipeline("translation", model="Helsinki-NLP/opus-mt-mul-en")
            model_cache[model_key] = {"pipeline": translator}
            return model_cache[model_key]

def translate_batch_marian(texts, src_lang, tgt_lang="en"):
    """Translate list of texts preserving order"""
    try:
        model_info = get_marian_model(src_lang, tgt_lang)
        if "pipeline" in model_info:
            return [model_info["pipeline"](t, max_length=512)[0]["translation_text"] if t.strip() else "" for t in texts]

        tokenizer = model_info["tokenizer"]
        model = model_info["model"]
        non_empty_texts = [t for t in texts if t.strip()]
        if not non_empty_texts:
            return texts

        inputs = tokenizer(non_empty_texts, return_tensors="pt", padding=True, truncation=True, max_length=512)
        with torch.no_grad():
            outputs = model.generate(**inputs, max_length=512, num_beams=4, early_stopping=True)
        translated_texts = tokenizer.batch_decode(outputs, skip_special_tokens=True)

        # Map back
        result = []
        idx = 0
        for t in texts:
            if t.strip():
                result.append(translated_texts[idx])
                idx += 1
            else:
                result.append("")
        return result
    except Exception as e:
        print(f"[Translation Error]: {e}")
        return [f"[Error: {e}]" for _ in texts]

def translate_text_preserve_format(text, src_lang, tgt_lang="en"):
    if src_lang == tgt_lang or not text.strip():
        return text
    lines = text.split("\n")
    translation_groups = [[line] for line in lines if line.strip()] or [[""]]
    translated = []
    for group in translation_groups:
        translated.extend(translate_batch_marian(group, src_lang, tgt_lang))
    return "\n".join(translated)

# ------------------------------
# File Handling
# ------------------------------
def read_file_content(filepath):
    ext = filepath.rsplit(".", 1)[1].lower()
    try:
        if ext == "txt":
            for enc in ["utf-8", "utf-8-sig", "latin-1", "cp1252"]:
                try:
                    with open(filepath, "r", encoding=enc) as f:
                        return f.read()
                except: continue
            raise ValueError("Cannot read text file")
        elif ext == "docx":
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs])
        elif ext == "pdf":
            reader = PdfReader(filepath)
            return "\n".join([p.extract_text() or "" for p in reader.pages])
        else:
            raise ValueError("Unsupported file")
    except Exception as e:
        raise ValueError(f"File read error: {e}")

def save_translated_content(content, filename, src_lang, tgt_lang, original_ext):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_name = filename.rsplit(".", 1)[0]
    if original_ext == "txt":
        out_name = f"{base_name}_{src_lang}_to_{tgt_lang}_{timestamp}.txt"
        out_path = os.path.join(OUTPUT_FOLDER, out_name)
        with open(out_path, "w", encoding="utf-8") as f: f.write(content)
        return out_path, out_name
    elif original_ext == "docx":
        input_path = os.path.join(UPLOAD_FOLDER, filename)
        doc_in = Document(input_path)
        doc_out = Document()
        out_name = f"{base_name}_{src_lang}_to_{tgt_lang}_{timestamp}.docx"
        out_path = os.path.join(OUTPUT_FOLDER, out_name)
        for para in doc_in.paragraphs:
            para_out = doc_out.add_paragraph()
            for run in para.runs:
                text = run.text.strip()
                if text:
                    translated = translate_text_preserve_format(text, src_lang, tgt_lang)
                    new_run = para_out.add_run(translated)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    new_run.font.size = run.font.size
                    new_run.font.name = run.font.name
                    if run.font.color and run.font.color.rgb:
                        new_run.font.color.rgb = RGBColor(run.font.color.rgb[0], run.font.color.rgb[1], run.font.color.rgb[2])
        doc_out.save(out_path)
        return out_path, out_name
    else:
        out_name = f"{base_name}_{src_lang}_to_{tgt_lang}_{timestamp}.docx"
        out_path = os.path.join(OUTPUT_FOLDER, out_name)
        doc = Document()
        doc.add_paragraph(content)
        doc.save(out_path)
        return out_path, out_name

# ------------------------------
# Flask Routes
# ------------------------------
@app.route("/", methods=["GET", "POST"])
def home():
    if request.method == "POST":
        input_type = request.form.get("input_type")
        tgt_lang = "en"
        if input_type == "file":
            if "file" not in request.files:
                flash("No file uploaded.", "error")
                return redirect(url_for("home"))
            file = request.files["file"]
            if file.filename == "":
                flash("No file selected.", "error")
                return redirect(url_for("home"))
            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                filepath = os.path.join(UPLOAD_FOLDER, filename)
                file.save(filepath)
                try:
                    start = datetime.now()
                    content = read_file_content(filepath)
                    if not content.strip():
                        flash("Uploaded file is empty.", "error")
                        os.remove(filepath)
                        return redirect(url_for("home"))

                    # Detect language
                    src_lang = detect_language_with_global_dialects(content)

                    if src_lang.split("-")[0] == tgt_lang:
                        translated = content
                        flash("File is already in English.", "success")
                    else:
                        translated = translate_text_preserve_format(content, src_lang, tgt_lang)
                        flash(f"File translated from {src_lang} to English!", "success")

                    original_ext = filename.rsplit(".", 1)[1].lower()
                    output_path, output_name = save_translated_content(translated, filename, src_lang, tgt_lang, original_ext)
                    os.remove(filepath)
                    end = datetime.now()
                    processing_time = (end - start).total_seconds()

                    return render_template_string(
                        HTML_TEMPLATE,
                        translated_text=translated[:2000]+"..." if len(translated)>2000 else translated,
                        download_url=url_for("download_file", filename=output_name),
                        detected_lang=src_lang,
                        processing_info=f"Translation completed in {processing_time:.2f}s. Language: {src_lang} → EN"
                    )
                except Exception as e:
                    if os.path.exists(filepath): os.remove(filepath)
                    flash(f"Error: {e}", "error")
                    return redirect(url_for("home"))
            else:
                flash("Invalid file type.", "error")
                return redirect(url_for("home"))
    return render_template_string(HTML_TEMPLATE)

@app.route("/ajax_translate", methods=["POST"])
def ajax_translate():
    try:
        data = request.get_json()
        text = data.get("text", "").strip()
        if not text:
            return jsonify({"translation": "", "new_part": "", "detected_lang": ""})
        src_lang = detect_language_with_global_dialects(text)
        if src_lang.split("-")[0] == "en":
            return jsonify({"translation": text, "new_part": text, "detected_lang": src_lang})
        translated = translate_text_preserve_format(text, src_lang, "en")
        return jsonify({"translation": translated, "new_part": translated, "detected_lang": src_lang})
    except Exception as e:
        return jsonify({"translation": "", "new_part": "", "error": str(e)})

@app.route("/download/<filename>")
def download_file(filename):
    try:
        return send_file(os.path.join(OUTPUT_FOLDER, filename), as_attachment=True, download_name=filename)
    except Exception as e:
        flash(f"Error downloading file: {e}", "error")
        return redirect(url_for("home"))

# ------------------------------
# Run Flask
# ------------------------------
if __name__ == "__main__":
    print("🚀 Starting Advanced AI Translator...")
    app.run(debug=True)

